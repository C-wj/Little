import os
import io
import tempfile
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from werkzeug.utils import secure_filename
from docx import Document
import markdownify
import zipfile
from datetime import datetime
from PIL import Image

app = Flask(__name__)
CORS(app)

# 配置
UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'docx', 'doc'}
ALLOWED_IMAGE_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif', 'webp', 'bmp'}
MAX_CONTENT_LENGTH = 16 * 1024 * 1024  # 16MB

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = MAX_CONTENT_LENGTH

# 确保上传目录存在
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


def allowed_file(filename):
    """检查文件扩展名是否允许"""
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def allowed_image_file(filename):
    """检查图片文件扩展名是否允许"""
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_IMAGE_EXTENSIONS


def compress_image_to_size(image_data, target_size_kb, output_format='JPEG'):
    """
    将图片压缩到指定大小（KB）
    使用二分法调整质量来逼近目标大小
    """
    target_size_bytes = target_size_kb * 1024
    
    # 打开图片
    img = Image.open(io.BytesIO(image_data))
    
    # 如果是 RGBA 模式且要输出为 JPEG，转换为 RGB
    if img.mode == 'RGBA' and output_format.upper() == 'JPEG':
        # 创建白色背景
        background = Image.new('RGB', img.size, (255, 255, 255))
        background.paste(img, mask=img.split()[3])  # 使用 alpha 通道作为 mask
        img = background
    elif img.mode != 'RGB' and output_format.upper() == 'JPEG':
        img = img.convert('RGB')
    
    # 先检查原图大小
    original_buffer = io.BytesIO()
    if output_format.upper() == 'PNG':
        img.save(original_buffer, format='PNG', optimize=True)
    else:
        img.save(original_buffer, format='JPEG', quality=95)
    original_size = original_buffer.tell()
    
    # 如果原图已经小于目标大小，直接返回
    if original_size <= target_size_bytes:
        original_buffer.seek(0)
        return original_buffer.getvalue(), original_size, 100, False
    
    # 使用二分法查找合适的质量值
    min_quality = 5
    max_quality = 95
    best_result = None
    best_size = float('inf')
    best_quality = min_quality
    
    # 同时考虑缩放图片
    scale_factor = 1.0
    current_img = img.copy()
    
    for _ in range(20):  # 最多迭代20次
        quality = (min_quality + max_quality) // 2
        
        buffer = io.BytesIO()
        if output_format.upper() == 'PNG':
            # PNG 使用压缩级别而不是质量
            current_img.save(buffer, format='PNG', optimize=True)
        else:
            current_img.save(buffer, format='JPEG', quality=quality, optimize=True)
        
        current_size = buffer.tell()
        
        if current_size <= target_size_bytes:
            if current_size > best_size * 0.5 or best_result is None:  # 选择更接近目标的
                best_result = buffer.getvalue()
                best_size = current_size
                best_quality = quality
            min_quality = quality + 1
        else:
            max_quality = quality - 1
        
        if min_quality > max_quality:
            break
    
    # 如果通过质量调整还是无法达到目标大小，需要缩放图片
    if best_size > target_size_bytes:
        # 计算需要的缩放比例
        scale_factor = (target_size_bytes / best_size) ** 0.5
        scale_factor = max(0.1, scale_factor)  # 最小缩放到 10%
        
        new_width = int(img.width * scale_factor)
        new_height = int(img.height * scale_factor)
        
        if new_width > 0 and new_height > 0:
            current_img = img.resize((new_width, new_height), Image.Resampling.LANCZOS)
            
            # 再次用二分法调整质量
            min_quality = 5
            max_quality = 95
            
            for _ in range(15):
                quality = (min_quality + max_quality) // 2
                
                buffer = io.BytesIO()
                if output_format.upper() == 'PNG':
                    current_img.save(buffer, format='PNG', optimize=True)
                else:
                    current_img.save(buffer, format='JPEG', quality=quality, optimize=True)
                
                current_size = buffer.tell()
                
                if current_size <= target_size_bytes:
                    best_result = buffer.getvalue()
                    best_size = current_size
                    best_quality = quality
                    min_quality = quality + 1
                else:
                    max_quality = quality - 1
                
                if min_quality > max_quality:
                    break
    
    if best_result is None:
        # 如果还是没有结果，使用最低质量
        buffer = io.BytesIO()
        if output_format.upper() == 'PNG':
            current_img.save(buffer, format='PNG', optimize=True)
        else:
            current_img.save(buffer, format='JPEG', quality=5, optimize=True)
        best_result = buffer.getvalue()
        best_size = buffer.tell()
        best_quality = 5
    
    resized = scale_factor < 1.0
    return best_result, best_size, best_quality, resized


def docx_to_markdown(docx_path, config=None):
    """将docx文件转换为markdown格式"""
    if config is None:
        config = get_default_config()
    
    try:
        # 读取docx文件
        doc = Document(docx_path)
        
        markdown_content = []
        paragraph_index = 0
        
        # 按顺序处理文档中的所有元素
        for element in doc.element.body:
            if element.tag.endswith('p'):  # 段落
                paragraph = None
                for p in doc.paragraphs:
                    if p._element == element:
                        paragraph = p
                        break
                
                if paragraph and paragraph.text.strip():
                    md_text = convert_paragraph_to_markdown(paragraph, paragraph_index, config)
                    if md_text:
                        markdown_content.append(md_text)
                    paragraph_index += 1
            
            elif element.tag.endswith('tbl'):  # 表格
                table = None
                for t in doc.tables:
                    if t._element == element:
                        table = t
                        break
                
                if table:
                    table_md = convert_table_to_markdown(table)
                    if table_md:
                        markdown_content.extend(table_md)
                        markdown_content.append('')  # 表格后添加空行
        
        # 清理空行并返回
        result = []
        for line in markdown_content:
            if line.strip() or (result and result[-1].strip()):
                result.append(line)
        
        return '\n'.join(result)
        
    except Exception as e:
        raise Exception(f"转换过程中出现错误: {str(e)}")


def get_default_config():
    """获取默认配置"""
    return {
        'treat_first_page_as_metadata': True,  # 第一页作为元数据处理
        'metadata_paragraph_limit': 10,  # 前N个段落可能是元数据
        'auto_detect_headers': True,  # 自动检测标题
        'exclude_table_headers_as_titles': True,  # 排除表格标题作为文档标题
        'min_heading_length': 2,  # 标题最小长度
        'max_heading_length': 80,  # 标题最大长度
    }


def convert_paragraph_to_markdown(paragraph, paragraph_index=0, config=None):
    """将段落转换为markdown格式"""
    if not paragraph.text.strip():
        return ''
    
    if config is None:
        config = get_default_config()
    
    # 检查是否为标题
    style_name = paragraph.style.name.lower()
    text = paragraph.text.strip()
    
    # 更准确的标题识别
    heading_patterns = {
        'heading 1': '#',
        'heading1': '#',
        '标题 1': '#',
        'heading 2': '##',
        'heading2': '##', 
        '标题 2': '##',
        'heading 3': '###',
        'heading3': '###',
        '标题 3': '###',
        'heading 4': '####',
        'heading4': '####',
        '标题 4': '####',
        'heading 5': '#####',
        'heading5': '#####',
        '标题 5': '#####',
        'heading 6': '######',
        'heading6': '######',
        '标题 6': '######'
    }
    
    # 检查是否匹配标题样式
    for pattern, prefix in heading_patterns.items():
        if pattern in style_name:
            return f'{prefix} {text}'
    
    # 对于没有使用标题样式的文档，尝试根据格式和内容判断标题
    if config['auto_detect_headers'] and is_likely_heading(paragraph, paragraph_index, config):
        heading_level = detect_heading_level(paragraph)
        prefix = '#' * heading_level
        return f'{prefix} {text}'
    
    # 如果不是标题，处理文本格式
    return format_text_runs(paragraph)


def is_likely_heading(paragraph, paragraph_index=0, config=None):
    """判断段落是否可能是标题"""
    if config is None:
        config = get_default_config()
        
    text = paragraph.text.strip()
    
    # 空文本不是标题
    if not text:
        return False
    
    # 检查长度限制
    if len(text) < config['min_heading_length'] or len(text) > config['max_heading_length']:
        return False
    
    # 前几个段落可能是文档元数据，需要更谨慎判断
    is_early_paragraph = paragraph_index < config['metadata_paragraph_limit']
    
    # 检查是否整个段落都是粗体
    if paragraph.runs:
        all_bold = True
        has_text = False
        for run in paragraph.runs:
            if run.text.strip():  # 忽略空白run
                has_text = True
                if not run.bold:
                    all_bold = False
                    break
        
        if not has_text:
            return False
            
        # 如果整个段落都是粗体，可能是标题
        if all_bold:
            # 排除明显的表格标题
            if config['exclude_table_headers_as_titles']:
                table_keywords = ['参数名称', '必须', '类型', '描述信息', 'retCode', 'exception']
                if any(keyword in text for keyword in table_keywords):
                    return False
            
            # 对于前几个段落，使用更严格的标准
            if is_early_paragraph:
                return is_likely_document_title(text, paragraph_index)
            else:
                # 后面的段落使用宽松的标准
                return (not text.endswith('。') and 
                       not text.endswith('.') and
                       not contains_complex_punctuation(text))
    
    return False


def is_likely_document_title(text, paragraph_index):
    """判断是否可能是文档标题（用于前几个段落的严格检查）"""
    # 明显的章节标识
    section_indicators = ['概述', '配置', '功能操作', '目录', '注意', '商标声明']
    if any(indicator in text for indicator in section_indicators):
        return True
    
    # 章节号格式
    if len(text.split()) >= 2:
        first_part = text.split()[0]
        if first_part.isdigit() and int(first_part) <= 20:
            return True
    
    # 数字编号格式 (如 "2.1", "3.2.1")
    if '.' in text and len(text) <= 15:
        parts = text.split('.')
        if len(parts) <= 3 and all(part.strip().isdigit() for part in parts if part.strip()):
            return True
    
    # 第一个段落如果很短且像标题，可能是文档标题
    if paragraph_index == 0 and len(text) <= 30:
        return True
    
    # 其他情况下，前几个段落比较可能是元数据
    return False


def contains_complex_punctuation(text):
    """检查是否包含复杂标点符号（表明可能是正文而非标题）"""
    complex_punct = ['，', '：', '；', '（', '）', '"', '"', '、']
    return any(punct in text for punct in complex_punct)


def detect_heading_level(paragraph):
    """检测标题级别"""
    text = paragraph.text.strip()
    
    # 基于文本内容和位置的简单启发式规则
    
    # 一级标题：主要章节
    if any(keyword in text for keyword in ['概述', '配置', '功能操作']):
        return 1
    
    # 检查是否是章节号 (如 "1 概述", "2 配置")
    if len(text.split()) >= 2:
        first_part = text.split()[0]
        if first_part.isdigit() and int(first_part) <= 10:  # 主要章节编号
            return 1
    
    # 二级标题：小节号 (如 "2.1", "3.2")
    if '.' in text and len(text) <= 10:
        parts = text.split('.')
        if len(parts) == 2 and parts[0].isdigit() and parts[1].isdigit():
            return 2
    
    # 三级标题：更细的分节 (如 "2.1.1")
    if text.count('.') == 2:
        parts = text.split('.')
        if len(parts) == 3 and all(part.isdigit() for part in parts):
            return 3
    
    # 基于关键词判断标题级别
    level2_keywords = ['目录', '商标声明', '注意']
    if any(keyword in text for keyword in level2_keywords):
        return 2
    
    # 默认二级标题
    return 2


def format_text_runs(paragraph):
    """处理段落中的文本格式（粗体、斜体等）"""
    if not paragraph.runs:
        return paragraph.text.strip()
    
    formatted_parts = []
    
    for run in paragraph.runs:
        text = run.text
        if not text:
            continue
            
        # 应用格式
        if run.bold and run.italic:
            text = f'***{text}***'
        elif run.bold:
            text = f'**{text}**'
        elif run.italic:
            text = f'*{text}*'
        
        formatted_parts.append(text)
    
    return ''.join(formatted_parts).strip()


def convert_table_to_markdown(table):
    """将表格转换为markdown格式"""
    if not table.rows:
        return []
    
    table_md = []
    
    for i, row in enumerate(table.rows):
        row_data = []
        for cell in row.cells:
            # 处理单元格内的文本，移除换行符
            cell_text = cell.text.strip().replace('\n', ' ').replace('\r', ' ')
            # 转义管道符号以避免表格格式冲突
            cell_text = cell_text.replace('|', '\\|')
            row_data.append(cell_text)
        
        # 生成表格行
        table_row = '| ' + ' | '.join(row_data) + ' |'
        table_md.append(table_row)
        
        # 如果是第一行（表头），添加分隔符
        if i == 0:
            separator = '|' + ' --- |' * len(row_data)
            table_md.append(separator)
    
    return table_md


def get_conversion_config(request):
    """从请求中获取转换配置"""
    config = get_default_config()
    
    # 从表单数据中获取配置选项
    if 'treat_first_page_as_metadata' in request.form:
        config['treat_first_page_as_metadata'] = request.form.get('treat_first_page_as_metadata', 'true').lower() == 'true'
    
    if 'metadata_paragraph_limit' in request.form:
        try:
            config['metadata_paragraph_limit'] = int(request.form.get('metadata_paragraph_limit', '10'))
        except ValueError:
            pass
    
    if 'auto_detect_headers' in request.form:
        config['auto_detect_headers'] = request.form.get('auto_detect_headers', 'true').lower() == 'true'
    
    if 'exclude_table_headers_as_titles' in request.form:
        config['exclude_table_headers_as_titles'] = request.form.get('exclude_table_headers_as_titles', 'true').lower() == 'true'
    
    if 'min_heading_length' in request.form:
        try:
            config['min_heading_length'] = int(request.form.get('min_heading_length', '2'))
        except ValueError:
            pass
    
    if 'max_heading_length' in request.form:
        try:
            config['max_heading_length'] = int(request.form.get('max_heading_length', '80'))
        except ValueError:
            pass
    
    return config


@app.route('/')
def index():
    """主页"""
    return '''
    <!DOCTYPE html>
    <html lang="zh-CN">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>Little 小工具集</title>
        <style>
            * {
                box-sizing: border-box;
                margin: 0;
                padding: 0;
            }
            
            body {
                font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
                background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                min-height: 100vh;
                display: flex;
                align-items: center;
                justify-content: center;
                padding: 20px;
            }
            
            .container {
                background: white;
                border-radius: 20px;
                padding: 40px;
                box-shadow: 0 20px 40px rgba(0,0,0,0.1);
                max-width: 650px;
                width: 100%;
                text-align: center;
            }
            
            h1 {
                color: #333;
                margin-bottom: 20px;
                font-size: 2.2em;
                background: linear-gradient(135deg, #667eea, #764ba2);
                -webkit-background-clip: text;
                -webkit-text-fill-color: transparent;
                background-clip: text;
            }
            
            .tabs {
                display: flex;
                margin-bottom: 30px;
                border-radius: 12px;
                overflow: hidden;
                background: #f0f3ff;
            }
            
            .tab {
                flex: 1;
                padding: 15px 20px;
                cursor: pointer;
                font-weight: 600;
                color: #667eea;
                transition: all 0.3s ease;
                border: none;
                background: transparent;
                font-size: 1em;
            }
            
            .tab:hover {
                background: rgba(102, 126, 234, 0.1);
            }
            
            .tab.active {
                background: linear-gradient(135deg, #667eea, #764ba2);
                color: white;
            }
            
            .tab-content {
                display: none;
            }
            
            .tab-content.active {
                display: block;
            }
            
            .upload-area {
                border: 3px dashed #667eea;
                border-radius: 15px;
                padding: 40px 20px;
                margin: 20px 0;
                transition: all 0.3s ease;
                cursor: pointer;
                background: #f8f9ff;
            }
            
            .upload-area:hover {
                border-color: #764ba2;
                background: #f0f3ff;
                transform: translateY(-2px);
            }
            
            .upload-area.dragover {
                border-color: #764ba2;
                background: #e8ecff;
                transform: scale(1.02);
            }
            
            .upload-icon {
                font-size: 3em;
                color: #667eea;
                margin-bottom: 15px;
            }
            
            .upload-text {
                color: #666;
                font-size: 1.1em;
                margin-bottom: 10px;
            }
            
            .file-input {
                display: none;
            }
            
            .btn {
                background: linear-gradient(135deg, #667eea, #764ba2);
                color: white;
                border: none;
                padding: 12px 30px;
                border-radius: 25px;
                font-size: 1em;
                cursor: pointer;
                transition: all 0.3s ease;
                margin: 10px;
                text-decoration: none;
                display: inline-block;
            }
            
            .btn:hover {
                transform: translateY(-2px);
                box-shadow: 0 10px 20px rgba(102, 126, 234, 0.3);
            }
            
            .btn:disabled {
                opacity: 0.6;
                cursor: not-allowed;
            }
            
            .progress {
                margin: 20px 0;
                display: none;
            }
            
            .progress-bar {
                width: 100%;
                height: 8px;
                background: #f0f0f0;
                border-radius: 4px;
                overflow: hidden;
            }
            
            .progress-fill {
                height: 100%;
                background: linear-gradient(135deg, #667eea, #764ba2);
                width: 0%;
                transition: width 0.3s ease;
            }
            
            .result {
                margin-top: 20px;
                padding: 20px;
                border-radius: 10px;
                display: none;
            }
            
            .result.success {
                background: #d4edda;
                border: 1px solid #c3e6cb;
                color: #155724;
            }
            
            .result.error {
                background: #f8d7da;
                border: 1px solid #f5c6cb;
                color: #721c24;
            }
            
            .config-panel {
                background: #f8f9ff;
                border-radius: 15px;
                padding: 20px;
                margin: 20px 0;
                border: 2px solid #e8ecff;
            }
            
            .config-header {
                color: #667eea;
                font-weight: bold;
                margin-bottom: 15px;
                cursor: pointer;
                display: flex;
                align-items: center;
                justify-content: space-between;
            }
            
            .config-content {
                display: none;
            }
            
            .config-content.show {
                display: block;
            }
            
            .config-row {
                display: flex;
                align-items: center;
                margin-bottom: 10px;
                flex-wrap: wrap;
                text-align: left;
            }
            
            .config-label {
                flex: 1;
                min-width: 200px;
                color: #555;
                font-size: 0.9em;
            }
            
            .config-input {
                flex: 0 0 auto;
                margin-left: 10px;
            }
            
            .config-input input[type="checkbox"] {
                transform: scale(1.2);
            }
            
            .config-input input[type="number"] {
                width: 100px;
                padding: 8px;
                border: 1px solid #ddd;
                border-radius: 8px;
                font-size: 0.95em;
            }
            
            .config-input select {
                width: 100px;
                padding: 8px;
                border: 1px solid #ddd;
                border-radius: 8px;
                font-size: 0.95em;
                background: white;
            }
            
            .size-setting {
                background: #f8f9ff;
                border-radius: 15px;
                padding: 20px;
                margin: 20px 0;
                border: 2px solid #e8ecff;
            }
            
            .size-setting-title {
                color: #667eea;
                font-weight: bold;
                margin-bottom: 15px;
                text-align: left;
            }
            
            .size-input-group {
                display: flex;
                align-items: center;
                justify-content: center;
                gap: 15px;
                flex-wrap: wrap;
            }
            
            .size-input-group label {
                color: #555;
                font-size: 0.95em;
            }
            
            .size-input-group input {
                width: 120px;
                padding: 10px;
                border: 2px solid #e8ecff;
                border-radius: 10px;
                font-size: 1em;
                text-align: center;
            }
            
            .size-input-group input:focus {
                outline: none;
                border-color: #667eea;
            }
            
            .size-input-group span {
                color: #667eea;
                font-weight: 600;
            }
            
            .compression-info {
                margin-top: 15px;
                padding: 15px;
                background: #e8f4fd;
                border-radius: 10px;
                font-size: 0.9em;
                color: #0c5460;
                text-align: left;
            }
            
            .compression-info strong {
                color: #667eea;
            }
            
            .preview-image {
                max-width: 100%;
                max-height: 200px;
                border-radius: 10px;
                margin: 15px 0;
                box-shadow: 0 4px 15px rgba(0,0,0,0.1);
            }
            
            @media (max-width: 600px) {
                .container {
                    margin: 10px;
                    padding: 30px 20px;
                }
                
                h1 {
                    font-size: 1.8em;
                }
                
                .tabs {
                    flex-direction: column;
                }
                
                .config-row {
                    flex-direction: column;
                    align-items: flex-start;
                }
                
                .config-input {
                    margin-left: 0;
                    margin-top: 5px;
                }
                
                .size-input-group {
                    flex-direction: column;
                }
            }
        </style>
    </head>
    <body>
        <div class="container">
            <h1>🛠️ Little 小工具集</h1>
            
            <div class="tabs">
                <button class="tab active" onclick="switchTab('word')">📝 Word转Markdown</button>
                <button class="tab" onclick="switchTab('image')">🖼️ 图片压缩</button>
            </div>
            
            <!-- Word转Markdown -->
            <div id="wordTab" class="tab-content active">
                <p style="color: #666; margin-bottom: 20px;">轻松将Word文档转换为Markdown格式</p>
                
                <div class="upload-area" id="wordUploadArea">
                    <div class="upload-icon">📁</div>
                    <div class="upload-text">点击或拖拽Word文档到这里</div>
                    <div style="color: #999; font-size: 0.9em;">支持 .docx 和 .doc 格式，最大16MB</div>
                    <input type="file" id="wordFileInput" class="file-input" accept=".docx,.doc" />
                </div>
                
                <div class="config-panel">
                    <div class="config-header" onclick="toggleConfig('word')">
                        <span>⚙️ 转换设置</span>
                        <span id="wordConfigToggle">▼</span>
                    </div>
                    <div class="config-content" id="wordConfigContent">
                        <div class="config-row">
                            <div class="config-label">智能标题检测</div>
                            <div class="config-input">
                                <input type="checkbox" id="autoDetectHeaders" checked>
                            </div>
                        </div>
                        <div class="config-row">
                            <div class="config-label">排除表格标题</div>
                            <div class="config-input">
                                <input type="checkbox" id="excludeTableHeaders" checked>
                            </div>
                        </div>
                        <div class="config-row">
                            <div class="config-label">元数据段落限制（前N段作为元数据）</div>
                            <div class="config-input">
                                <input type="number" id="metadataLimit" value="10" min="0" max="50">
                            </div>
                        </div>
                        <div class="config-row">
                            <div class="config-label">标题最小长度</div>
                            <div class="config-input">
                                <input type="number" id="minHeadingLength" value="2" min="1" max="20">
                            </div>
                        </div>
                        <div class="config-row">
                            <div class="config-label">标题最大长度</div>
                            <div class="config-input">
                                <input type="number" id="maxHeadingLength" value="80" min="10" max="200">
                            </div>
                        </div>
                    </div>
                </div>
                
                <div class="progress" id="wordProgress">
                    <div class="progress-bar">
                        <div class="progress-fill" id="wordProgressFill"></div>
                    </div>
                    <div style="margin-top: 10px; color: #666;">转换中...</div>
                </div>
                
                <div class="result" id="wordResult">
                    <div id="wordResultMessage"></div>
                    <div id="wordDownloadLink" style="margin-top: 15px;"></div>
                </div>
            </div>
            
            <!-- 图片压缩 -->
            <div id="imageTab" class="tab-content">
                <p style="color: #666; margin-bottom: 20px;">将图片压缩到指定文件大小</p>
                
                <div class="upload-area" id="imageUploadArea">
                    <div class="upload-icon">🖼️</div>
                    <div class="upload-text">点击或拖拽图片到这里</div>
                    <div style="color: #999; font-size: 0.9em;">支持 PNG、JPG、JPEG、GIF、WebP、BMP 格式</div>
                    <input type="file" id="imageFileInput" class="file-input" accept=".png,.jpg,.jpeg,.gif,.webp,.bmp" />
                </div>
                
                <img id="imagePreview" class="preview-image" style="display: none;" />
                
                <div class="size-setting">
                    <div class="size-setting-title">📐 压缩设置</div>
                    <div class="size-input-group">
                        <label for="targetSize">目标大小：</label>
                        <input type="number" id="targetSize" value="500" min="1" max="10240" />
                        <span>KB</span>
                    </div>
                    <div class="config-row" style="margin-top: 15px;">
                        <div class="config-label">输出格式</div>
                        <div class="config-input">
                            <select id="outputFormat">
                                <option value="JPEG" selected>JPEG</option>
                                <option value="PNG">PNG</option>
                                <option value="WEBP">WebP</option>
                            </select>
                        </div>
                    </div>
                    <div class="compression-info" id="originalSizeInfo" style="display: none;">
                        <strong>原图大小：</strong><span id="originalSize">-</span>
                    </div>
                </div>
                
                <div class="progress" id="imageProgress">
                    <div class="progress-bar">
                        <div class="progress-fill" id="imageProgressFill"></div>
                    </div>
                    <div style="margin-top: 10px; color: #666;">压缩中...</div>
                </div>
                
                <div class="result" id="imageResult">
                    <div id="imageResultMessage"></div>
                    <div id="imageDownloadLink" style="margin-top: 15px;"></div>
                </div>
            </div>
        </div>

        <script>
            // 标签切换
            function switchTab(tab) {
                document.querySelectorAll('.tab').forEach(t => t.classList.remove('active'));
                document.querySelectorAll('.tab-content').forEach(c => c.classList.remove('active'));
                
                if (tab === 'word') {
                    document.querySelector('.tab:nth-child(1)').classList.add('active');
                    document.getElementById('wordTab').classList.add('active');
                } else {
                    document.querySelector('.tab:nth-child(2)').classList.add('active');
                    document.getElementById('imageTab').classList.add('active');
                }
            }

            // 配置面板切换
            function toggleConfig(type) {
                const content = document.getElementById(type + 'ConfigContent');
                const toggle = document.getElementById(type + 'ConfigToggle');
                
                if (content.classList.contains('show')) {
                    content.classList.remove('show');
                    toggle.textContent = '▼';
                } else {
                    content.classList.add('show');
                    toggle.textContent = '▲';
                }
            }

            // ========== Word转Markdown ==========
            const wordUploadArea = document.getElementById('wordUploadArea');
            const wordFileInput = document.getElementById('wordFileInput');
            const wordProgress = document.getElementById('wordProgress');
            const wordProgressFill = document.getElementById('wordProgressFill');
            const wordResult = document.getElementById('wordResult');
            const wordResultMessage = document.getElementById('wordResultMessage');
            const wordDownloadLink = document.getElementById('wordDownloadLink');

            function getConversionConfig() {
                return {
                    auto_detect_headers: document.getElementById('autoDetectHeaders').checked,
                    exclude_table_headers_as_titles: document.getElementById('excludeTableHeaders').checked,
                    metadata_paragraph_limit: document.getElementById('metadataLimit').value,
                    min_heading_length: document.getElementById('minHeadingLength').value,
                    max_heading_length: document.getElementById('maxHeadingLength').value
                };
            }

            wordUploadArea.addEventListener('click', () => wordFileInput.click());
            
            wordUploadArea.addEventListener('dragover', (e) => {
                e.preventDefault();
                wordUploadArea.classList.add('dragover');
            });

            wordUploadArea.addEventListener('dragleave', () => {
                wordUploadArea.classList.remove('dragover');
            });

            wordUploadArea.addEventListener('drop', (e) => {
                e.preventDefault();
                wordUploadArea.classList.remove('dragover');
                if (e.dataTransfer.files.length > 0) {
                    handleWordFile(e.dataTransfer.files[0]);
                }
            });

            wordFileInput.addEventListener('change', (e) => {
                if (e.target.files.length > 0) {
                    handleWordFile(e.target.files[0]);
                }
            });

            function handleWordFile(file) {
                if (!file.name.endsWith('.docx') && !file.name.endsWith('.doc')) {
                    showWordResult('error', '请选择Word文档文件 (.docx 或 .doc)');
                    return;
                }
                if (file.size > 16 * 1024 * 1024) {
                    showWordResult('error', '文件大小超过16MB限制');
                    return;
                }
                uploadWordFile(file);
            }

            function uploadWordFile(file) {
                const formData = new FormData();
                formData.append('file', file);
                
                const config = getConversionConfig();
                for (const [key, value] of Object.entries(config)) {
                    formData.append(key, value);
                }

                wordProgress.style.display = 'block';
                wordResult.style.display = 'none';
                wordProgressFill.style.width = '0%';

                let progressValue = 0;
                const progressInterval = setInterval(() => {
                    progressValue += Math.random() * 20;
                    if (progressValue > 90) progressValue = 90;
                    wordProgressFill.style.width = progressValue + '%';
                }, 100);

                fetch('/convert', {
                    method: 'POST',
                    body: formData
                })
                .then(response => {
                    clearInterval(progressInterval);
                    wordProgressFill.style.width = '100%';
                    
                    if (response.ok) {
                        return response.blob();
                    } else {
                        return response.json().then(err => Promise.reject(err));
                    }
                })
                .then(blob => {
                    const url = window.URL.createObjectURL(blob);
                    const filename = file.name.replace(/\\.(docx?|doc)$/i, '.md');
                    
                    showWordResult('success', '转换成功！');
                    wordDownloadLink.innerHTML = `<a href="${url}" download="${filename}" class="btn">📥 下载Markdown文件</a>`;
                    
                    setTimeout(() => {
                        wordProgress.style.display = 'none';
                    }, 500);
                })
                .catch(error => {
                    clearInterval(progressInterval);
                    wordProgress.style.display = 'none';
                    showWordResult('error', error.message || '转换失败，请重试');
                });
            }

            function showWordResult(type, message) {
                wordResult.className = `result ${type}`;
                wordResult.style.display = 'block';
                wordResultMessage.textContent = message;
                if (type === 'error') {
                    wordDownloadLink.innerHTML = '';
                }
            }

            // ========== 图片压缩 ==========
            const imageUploadArea = document.getElementById('imageUploadArea');
            const imageFileInput = document.getElementById('imageFileInput');
            const imageProgress = document.getElementById('imageProgress');
            const imageProgressFill = document.getElementById('imageProgressFill');
            const imageResult = document.getElementById('imageResult');
            const imageResultMessage = document.getElementById('imageResultMessage');
            const imageDownloadLink = document.getElementById('imageDownloadLink');
            const imagePreview = document.getElementById('imagePreview');
            const originalSizeInfo = document.getElementById('originalSizeInfo');
            const originalSizeSpan = document.getElementById('originalSize');

            let selectedImageFile = null;

            imageUploadArea.addEventListener('click', () => imageFileInput.click());

            imageUploadArea.addEventListener('dragover', (e) => {
                e.preventDefault();
                imageUploadArea.classList.add('dragover');
            });

            imageUploadArea.addEventListener('dragleave', () => {
                imageUploadArea.classList.remove('dragover');
            });

            imageUploadArea.addEventListener('drop', (e) => {
                e.preventDefault();
                imageUploadArea.classList.remove('dragover');
                if (e.dataTransfer.files.length > 0) {
                    handleImageFile(e.dataTransfer.files[0]);
                }
            });

            imageFileInput.addEventListener('change', (e) => {
                if (e.target.files.length > 0) {
                    handleImageFile(e.target.files[0]);
                }
            });

            function formatFileSize(bytes) {
                if (bytes < 1024) return bytes + ' B';
                if (bytes < 1024 * 1024) return (bytes / 1024).toFixed(2) + ' KB';
                return (bytes / (1024 * 1024)).toFixed(2) + ' MB';
            }

            function handleImageFile(file) {
                const validExtensions = ['png', 'jpg', 'jpeg', 'gif', 'webp', 'bmp'];
                const ext = file.name.split('.').pop().toLowerCase();
                
                if (!validExtensions.includes(ext)) {
                    showImageResult('error', '请选择有效的图片文件');
                    return;
                }
                
                if (file.size > 16 * 1024 * 1024) {
                    showImageResult('error', '文件大小超过16MB限制');
                    return;
                }

                selectedImageFile = file;
                
                // 显示预览
                const reader = new FileReader();
                reader.onload = (e) => {
                    imagePreview.src = e.target.result;
                    imagePreview.style.display = 'block';
                };
                reader.readAsDataURL(file);
                
                // 显示原图大小
                originalSizeSpan.textContent = formatFileSize(file.size);
                originalSizeInfo.style.display = 'block';
                
                // 清除之前的结果
                imageResult.style.display = 'none';
                imageDownloadLink.innerHTML = '';
                
                // 自动开始压缩
                compressImage();
            }

            function compressImage() {
                if (!selectedImageFile) {
                    showImageResult('error', '请先选择图片');
                    return;
                }

                const targetSize = document.getElementById('targetSize').value;
                const outputFormat = document.getElementById('outputFormat').value;

                const formData = new FormData();
                formData.append('file', selectedImageFile);
                formData.append('target_size', targetSize);
                formData.append('output_format', outputFormat);

                imageProgress.style.display = 'block';
                imageResult.style.display = 'none';
                imageProgressFill.style.width = '0%';

                let progressValue = 0;
                const progressInterval = setInterval(() => {
                    progressValue += Math.random() * 15;
                    if (progressValue > 90) progressValue = 90;
                    imageProgressFill.style.width = progressValue + '%';
                }, 100);

                fetch('/compress-image', {
                    method: 'POST',
                    body: formData
                })
                .then(response => {
                    clearInterval(progressInterval);
                    imageProgressFill.style.width = '100%';
                    
                    if (response.ok) {
                        const compressedSize = response.headers.get('X-Compressed-Size');
                        const quality = response.headers.get('X-Compression-Quality');
                        const resized = response.headers.get('X-Image-Resized');
                        
                        return response.blob().then(blob => ({
                            blob,
                            compressedSize,
                            quality,
                            resized
                        }));
                    } else {
                        return response.json().then(err => Promise.reject(err));
                    }
                })
                .then(({blob, compressedSize, quality, resized}) => {
                    const url = window.URL.createObjectURL(blob);
                    const ext = outputFormat.toLowerCase() === 'jpeg' ? 'jpg' : outputFormat.toLowerCase();
                    const baseName = selectedImageFile.name.replace(/\\.[^.]+$/, '');
                    const filename = `${baseName}_compressed.${ext}`;
                    
                    let message = `压缩成功！压缩后大小：${formatFileSize(parseInt(compressedSize))}`;
                    if (resized === 'true') {
                        message += '（已自动调整尺寸）';
                    }
                    
                    showImageResult('success', message);
                    imageDownloadLink.innerHTML = `<a href="${url}" download="${filename}" class="btn">📥 下载压缩后的图片</a>`;
                    
                    setTimeout(() => {
                        imageProgress.style.display = 'none';
                    }, 500);
                })
                .catch(error => {
                    clearInterval(progressInterval);
                    imageProgress.style.display = 'none';
                    showImageResult('error', error.message || '压缩失败，请重试');
                });
            }

            function showImageResult(type, message) {
                imageResult.className = `result ${type}`;
                imageResult.style.display = 'block';
                imageResultMessage.textContent = message;
                if (type === 'error') {
                    imageDownloadLink.innerHTML = '';
                }
            }

            // 目标大小或格式改变时重新压缩
            document.getElementById('targetSize').addEventListener('change', () => {
                if (selectedImageFile) compressImage();
            });
            document.getElementById('outputFormat').addEventListener('change', () => {
                if (selectedImageFile) compressImage();
            });
        </script>
    </body>
    </html>
    '''


@app.route('/convert', methods=['POST'])
def convert_word_to_markdown():
    """转换Word文档为Markdown"""
    try:
        # 检查是否有文件上传
        if 'file' not in request.files:
            return jsonify({'error': '没有选择文件'}), 400
        
        file = request.files['file']
        
        # 检查文件名
        if file.filename == '':
            return jsonify({'error': '没有选择文件'}), 400
        
        # 检查文件类型
        if not allowed_file(file.filename):
            return jsonify({'error': '不支持的文件格式，请选择Word文档(.docx或.doc)'}), 400
        
        # 保存上传的文件
        filename = secure_filename(file.filename)
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        unique_filename = f"{timestamp}_{filename}"
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], unique_filename)
        file.save(filepath)
        
        try:
            # 获取转换配置
            config = get_conversion_config(request)
            
            # 转换为Markdown
            markdown_content = docx_to_markdown(filepath, config)
            
            # 创建临时文件用于下载
            with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False, encoding='utf-8') as temp_file:
                temp_file.write(markdown_content)
                temp_filepath = temp_file.name
            
            # 生成下载文件名
            base_name = os.path.splitext(filename)[0]
            download_filename = f"{base_name}.md"
            
            # 清理上传的文件
            os.remove(filepath)
            
            # 返回文件用于下载
            return send_file(
                temp_filepath,
                as_attachment=True,
                download_name=download_filename,
                mimetype='text/markdown'
            )
            
        except Exception as e:
            # 清理上传的文件
            if os.path.exists(filepath):
                os.remove(filepath)
            return jsonify({'error': str(e)}), 500
            
    except Exception as e:
        return jsonify({'error': f'处理请求时出现错误: {str(e)}'}), 500


@app.route('/compress-image', methods=['POST'])
def compress_image():
    """压缩图片到指定大小"""
    try:
        # 检查是否有文件上传
        if 'file' not in request.files:
            return jsonify({'error': '没有选择文件'}), 400
        
        file = request.files['file']
        
        # 检查文件名
        if file.filename == '':
            return jsonify({'error': '没有选择文件'}), 400
        
        # 检查文件类型
        if not allowed_image_file(file.filename):
            return jsonify({'error': '不支持的图片格式，请选择 PNG、JPG、JPEG、GIF、WebP 或 BMP 格式'}), 400
        
        # 获取目标大小（KB）
        target_size_kb = request.form.get('target_size', type=int, default=500)
        if target_size_kb < 1:
            return jsonify({'error': '目标大小必须大于 0'}), 400
        if target_size_kb > 10240:  # 最大 10MB
            return jsonify({'error': '目标大小不能超过 10MB'}), 400
        
        # 获取输出格式
        output_format = request.form.get('output_format', 'JPEG').upper()
        if output_format not in ['JPEG', 'PNG', 'WEBP']:
            output_format = 'JPEG'
        
        # 读取图片数据
        image_data = file.read()
        original_size = len(image_data)
        
        try:
            # 压缩图片
            compressed_data, final_size, quality, resized = compress_image_to_size(
                image_data, target_size_kb, output_format
            )
            
            # 确定输出文件扩展名
            ext_map = {'JPEG': 'jpg', 'PNG': 'png', 'WEBP': 'webp'}
            output_ext = ext_map.get(output_format, 'jpg')
            
            # 生成下载文件名
            base_name = os.path.splitext(secure_filename(file.filename))[0]
            download_filename = f"{base_name}_compressed.{output_ext}"
            
            # 创建临时文件
            with tempfile.NamedTemporaryFile(delete=False, suffix=f'.{output_ext}') as temp_file:
                temp_file.write(compressed_data)
                temp_filepath = temp_file.name
            
            # 设置响应头，包含压缩信息
            response = send_file(
                temp_filepath,
                as_attachment=True,
                download_name=download_filename,
                mimetype=f'image/{output_ext}'
            )
            
            # 添加自定义响应头
            response.headers['X-Original-Size'] = str(original_size)
            response.headers['X-Compressed-Size'] = str(final_size)
            response.headers['X-Compression-Quality'] = str(quality)
            response.headers['X-Image-Resized'] = str(resized).lower()
            
            return response
            
        except Exception as e:
            return jsonify({'error': f'压缩图片时出现错误: {str(e)}'}), 500
            
    except Exception as e:
        return jsonify({'error': f'处理请求时出现错误: {str(e)}'}), 500


@app.route('/health')
def health_check():
    """健康检查接口"""
    return jsonify({'status': 'ok', 'message': '小工具集运行正常'})


if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=8080)