#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
调试Word文档结构，分析标题和段落样式
"""

from docx import Document
import sys

def analyze_document_structure(docx_path):
    """分析Word文档的结构"""
    try:
        doc = Document(docx_path)
        print(f"分析文档: {docx_path}")
        print("=" * 50)
        
        # 分析前20个段落的样式
        for i, paragraph in enumerate(doc.paragraphs[:20]):
            if paragraph.text.strip():
                style_name = paragraph.style.name
                text_preview = paragraph.text.strip()[:50] + "..." if len(paragraph.text.strip()) > 50 else paragraph.text.strip()
                
                print(f"段落 {i+1}:")
                print(f"  样式名称: '{style_name}'")
                print(f"  文本内容: '{text_preview}'")
                
                # 检查是否有特殊格式
                if paragraph.runs:
                    formats = []
                    for run in paragraph.runs:
                        if run.bold:
                            formats.append("粗体")
                        if run.italic:
                            formats.append("斜体")
                    if formats:
                        print(f"  格式: {', '.join(formats)}")
                
                print()
        
        print("\n可用的样式列表:")
        print("=" * 30)
        styles = set()
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                styles.add(paragraph.style.name)
        
        for style in sorted(styles):
            print(f"- '{style}'")
            
    except Exception as e:
        print(f"分析出错: {e}")

if __name__ == '__main__':
    analyze_document_structure('VOS3000 Web第三方接口文档.docx')